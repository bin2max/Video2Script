import os
import json
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import config

class DocumentProcessor:
    def __init__(self):
        self.output_dir = config.OUTPUT_DIR
        self.ensure_output_dir()
    
    def ensure_output_dir(self):
        """确保输出目录存在"""
        if not os.path.exists(self.output_dir):
            os.makedirs(self.output_dir)
    
    def load_template(self, template_path):
        """加载Word文档模板"""
        try:
            if os.path.exists(template_path):
                doc = Document(template_path)
                return doc
            else:
                return self.create_default_template()
        except Exception as e:
            print(f"模板加载失败: {str(e)}")
            return self.create_default_template()
    
    def create_default_template(self):
        """创建默认模板"""
        doc = Document()
        
        # 设置页面边距
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
        
        # 添加标题
        title = doc.add_heading('技能操作培训脚本', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 添加基本信息
        doc.add_paragraph('培训日期：')
        doc.add_paragraph('培训讲师：')
        doc.add_paragraph('培训地点：')
        
        doc.add_paragraph()  # 空行
        
        return doc
    
    def apply_template_structure(self, doc, structured_content):
        """将结构化内容应用到文档模板"""
        # 清空现有内容（保留标题）
        for i in range(len(doc.paragraphs) - 1, 0, -1):
            p = doc.paragraphs[i]
            p._element.getparent().remove(p._element)
        
        # 添加标题
        title = doc.add_heading(structured_content['title'], 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 添加基本信息
        doc.add_paragraph(f'生成日期：{datetime.now().strftime("%Y年%m月%d日")}')
        doc.add_paragraph(f'生成时间：{datetime.now().strftime("%H:%M:%S")}')
        doc.add_paragraph('文档类型：自动生成的培训脚本')
        
        doc.add_paragraph()  # 空行
        
        # 添加各个章节
        for section in structured_content['sections']:
            # 添加章节标题
            doc.add_heading(section['title'], level=1)
            
            # 添加章节内容
            for content in section['content']:
                doc.add_paragraph(content)
            
            doc.add_paragraph()  # 章节间空行
        
        return doc
    
    def export_to_docx(self, structured_content, template_path=None, output_filename=None):
        """导出为Word文档"""
        try:
            # 加载模板
            if template_path and os.path.exists(template_path):
                doc = self.load_template(template_path)
            else:
                doc = self.create_default_template()
            
            # 应用内容结构
            doc = self.apply_template_structure(doc, structured_content)
            
            # 生成输出文件名
            if not output_filename:
                timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                output_filename = f"培训脚本_{timestamp}.docx"
            
            output_path = os.path.join(self.output_dir, output_filename)
            
            # 保存文档
            doc.save(output_path)
            
            print(f"文档已保存到: {output_path}")
            return output_path
            
        except Exception as e:
            raise Exception(f"文档导出失败: {str(e)}")
    
    def save_structured_content(self, structured_content, filename=None):
        """保存结构化内容为JSON文件（用于备份）"""
        try:
            if not filename:
                timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                filename = f"结构化内容_{timestamp}.json"
            
            filepath = os.path.join(self.output_dir, filename)
            
            with open(filepath, 'w', encoding='utf-8') as f:
                json.dump(structured_content, f, ensure_ascii=False, indent=2)
            
            print(f"结构化内容已保存到: {filepath}")
            return filepath
            
        except Exception as e:
            print(f"保存结构化内容失败: {str(e)}")
            return None
    
    def load_structured_content(self, filepath):
        """从JSON文件加载结构化内容"""
        try:
            with open(filepath, 'r', encoding='utf-8') as f:
                structured_content = json.load(f)
            return structured_content
        except Exception as e:
            print(f"加载结构化内容失败: {str(e)}")
            return None

class MockDocumentProcessor:
    """模拟文档处理器，用于测试"""
    def __init__(self):
        self.output_dir = config.OUTPUT_DIR
        self.ensure_output_dir()
    
    def ensure_output_dir(self):
        """确保输出目录存在"""
        if not os.path.exists(self.output_dir):
            os.makedirs(self.output_dir)
    
    def export_to_docx(self, structured_content, template_path=None, output_filename=None):
        """模拟导出Word文档"""
        try:
            # 生成输出文件名
            if not output_filename:
                timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                output_filename = f"培训脚本_{timestamp}.docx"
            
            output_path = os.path.join(self.output_dir, output_filename)
            
            # 模拟创建文档
            doc = Document()
            title = doc.add_heading(structured_content['title'], 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # 添加基本信息
            doc.add_paragraph(f'生成日期：{datetime.now().strftime("%Y年%m月%d日")}')
            doc.add_paragraph(f'生成时间：{datetime.now().strftime("%H:%M:%S")}')
            doc.add_paragraph('文档类型：自动生成的培训脚本')
            
            # 添加章节内容
            for section in structured_content['sections']:
                doc.add_heading(section['title'], level=1)
                for content in section['content']:
                    doc.add_paragraph(content)
                doc.add_paragraph()
            
            # 保存文档
            doc.save(output_path)
            
            print(f"文档已保存到: {output_path}")
            return output_path
            
        except Exception as e:
            print(f"文档导出失败: {str(e)}")
            return None 